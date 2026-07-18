"""Regression tests for issues found by the full project review."""

from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document as DocxDocument

from ccitecheck.recognition.arbitration import arbitrate_claim_candidates, build_claim_document
from ccitecheck.recognition.cases import extract_case_refs
from ccitecheck.recognition.service import build_indexes
from ccitecheck.recognition.statutes import extract_legal_sources
from ccitecheck.recognition.rules import extract_rule_candidates
from ccitecheck.domain.citation import (
    CaseCitationEntities,
    CaseRef,
    CaseReferenceType,
    ClaimCandidate,
    ClaimDocument,
    ClaimType,
)
from ccitecheck.infrastructure.database import (
    connect,
    find_current_article,
    init_db,
    normalize_article_key,
    upsert_article,
    upsert_law,
)
from ccitecheck.parsing.docx import parse_docx
from ccitecheck.domain.document import Anchor, Block, BlockType, Chunk, DocMeta, ParsedDocument
from ccitecheck.judgment.cases import verify_case_claims
from ccitecheck.domain.evidence import CaseLookupStatus


def _parsed(*texts: str) -> ParsedDocument:
    blocks = []
    anchors = []
    for index, value in enumerate(texts, start=1):
        block_id = f"b_{index:05d}"
        anchor_id = f"line{index:05d}"
        blocks.append(
            Block(
                block_id=block_id,
                type=BlockType.PARAGRAPH,
                text=value,
                body_order=index - 1,
                block_order=index - 1,
                para_index=index - 1,
                anchor_range=[anchor_id, anchor_id],
                sentence_anchors=[anchor_id],
            )
        )
        anchors.append(
            Anchor(
                anchor=anchor_id,
                text=value,
                block_id=block_id,
                para_index=index - 1,
                char_start=0,
                char_end=len(value),
            )
        )
    return ParsedDocument(
        doc_meta=DocMeta(source_file="review.docx", doc_hash="sha256:review"),
        blocks=blocks,
        anchors=anchors,
        chunks=[
            Chunk(
                chunk_id="c_00001",
                block_ids=[block.block_id for block in blocks],
                anchor_ids=[anchor.anchor for anchor in anchors],
                anchor_range=[anchors[0].anchor, anchors[-1].anchor],
            )
        ],
    )


def test_article_key_normalizes_chinese_arabic_and_zhi_suffix():
    assert normalize_article_key("第一百二十七条") == "127"
    assert normalize_article_key("第127条") == "127"
    assert normalize_article_key("第一百八十四条之一") == "184-1"


def test_chinese_article_key_is_found_by_arabic_citation(tmp_path: Path):
    db_path = tmp_path / "laws.sqlite"
    init_db(db_path)
    with connect(db_path) as conn:
        law_id = upsert_law(conn, {"title": "中华人民共和国民法典", "source_type": "law"})
        article_id = upsert_article(
            conn,
            law_id,
            {"article_no": "第一百二十七条", "text": "依法保护数据和网络虚拟财产。"},
        )
        conn.execute(
            "UPDATE articles SET article_key = ? WHERE id = ?",
            ("一百二十七", article_id),
        )
    with connect(db_path) as conn:
        article = find_current_article(conn, "民法典", "第127条")
    assert article is not None
    assert article["article_no"] == "第一百二十七条"


def test_mixed_statute_and_case_sentence_yields_both_candidates():
    doc = _parsed("依据《民法典》第五百七十七条及（2025）晋民再174号，应核对两类引用。")
    candidates = extract_rule_candidates(doc, build_indexes(doc))
    assert {candidate.claim_type for candidate in candidates} == {
        ClaimType.LEGAL_SOURCE_CLAIM,
        ClaimType.CASE_CITATION,
    }


def test_adjacent_table_law_name_and_article_cells_merge_into_one_claim():
    doc = _parsed("《中华人民共和国民法典》", "第127条规定数据权益。")
    for index, block in enumerate(doc.blocks):
        block.type = BlockType.TABLE_CELL
        block.table_index = 0
        block.row_index = 0
        block.cell_index = index
    candidates = extract_rule_candidates(doc, build_indexes(doc), include_cases=False)
    assert len(candidates) == 1
    assert candidates[0].anchor_ids == ["line00001", "line00002"]
    source = candidates[0].entities.legal_sources[0]
    assert source.title == "中华人民共和国民法典"
    assert source.articles[0].article == "第127条"
    assert source.resolution == "inherited"
    assert source.inherited_from_anchor == "line00001"

    claims = arbitrate_claim_candidates(candidates, doc)
    assert claims[0].source_locations[-1].cell_index == 1
    inherited_location = claims[0].entities.legal_sources[0].inherited_from_location
    assert inherited_location is not None
    assert inherited_location.cell_index == 0


def test_article_paragraphs_and_items_stop_at_next_article():
    sources = extract_legal_sources(
        "依据《民法典》第十二条第一款第一项；第十三条第二款第四项处理。"
    )
    first, second = sources[0].articles
    assert first.paragraphs == ["第一款"]
    assert first.items == ["第一项"]
    assert second.paragraphs == ["第二款"]
    assert second.items == ["第四项"]


def test_next_sentence_in_same_paragraph_inherits_unique_law_source(tmp_path: Path):
    text = (
        "《最高人民法院关于审理涉及驰名商标保护的民事纠纷案件应用法律若干问题的解释》"
        "第九条规定，相关公众容易产生混淆。"
        "第十条规定，人民法院应当综合考虑相关因素后作出裁判。"
    )
    path = tmp_path / "same-paragraph.docx"
    document = DocxDocument()
    document.add_paragraph(text)
    document.save(path)
    parsed = parse_docx(str(path))
    candidates = extract_rule_candidates(parsed, build_indexes(parsed), include_cases=False)
    assert len(candidates) == 2
    second_source = candidates[1].entities.legal_sources[0]
    assert second_source.title == "最高人民法院关于审理涉及驰名商标保护的民事纠纷案件应用法律若干问题的解释"
    assert second_source.articles[0].article == "第十条"
    assert second_source.resolution == "inherited"


def test_same_paragraph_partial_paragraph_and_item_inherit_unique_article(tmp_path: Path):
    text = (
        "《广告法》第二十八条第一款规定，虚假内容构成虚假广告。"
        "第二款第二项规定，商品性能等信息与实际情况不符的，为虚假广告。"
    )
    path = tmp_path / "partial-paragraph-item.docx"
    document = DocxDocument()
    document.add_paragraph(text)
    document.save(path)
    parsed = parse_docx(str(path))
    candidates = extract_rule_candidates(parsed, build_indexes(parsed), include_cases=False)
    assert len(candidates) == 2
    second_source = candidates[1].entities.legal_sources[0]
    second_article = second_source.articles[0]
    assert second_source.title == "广告法"
    assert second_source.resolution == "inherited"
    assert second_article.article == "第二十八条"
    assert second_article.paragraphs == ["第二款"]
    assert second_article.items == ["第二项"]


def test_partial_item_only_inherits_unique_parent_paragraph(tmp_path: Path):
    text = "《广告法》第二十八条第二款规定了五种情形。第二项规定，商品性能信息不得失实。"
    path = tmp_path / "partial-item.docx"
    document = DocxDocument()
    document.add_paragraph(text)
    document.save(path)
    parsed = parse_docx(str(path))
    candidates = extract_rule_candidates(parsed, build_indexes(parsed), include_cases=False)
    second_article = candidates[1].entities.legal_sources[0].articles[0]
    assert second_article.article == "第二十八条"
    assert second_article.paragraphs == ["第二款"]
    assert second_article.items == ["第二项"]


def test_partial_ref_does_not_inherit_ambiguous_parent_articles(tmp_path: Path):
    text = "《广告法》第二十八条、第二十九条规定了不同规则。第二款规定了具体条件。"
    path = tmp_path / "ambiguous-parent-articles.docx"
    document = DocxDocument()
    document.add_paragraph(text)
    document.save(path)
    parsed = parse_docx(str(path))
    candidates = extract_rule_candidates(parsed, build_indexes(parsed), include_cases=False)
    assert len(candidates) == 1


def test_partial_ref_rejects_ordinary_product_and_task_wording(tmp_path: Path):
    text = "《广告法》第二十八条规定了虚假广告。公司发布第二款产品，随后完成第三项任务。"
    path = tmp_path / "ordinary-ordinal-wording.docx"
    document = DocxDocument()
    document.add_paragraph(text)
    document.save(path)
    parsed = parse_docx(str(path))
    candidates = extract_rule_candidates(parsed, build_indexes(parsed), include_cases=False)
    assert len(candidates) == 1


def test_source_location_occurrence_counts_repeated_anchor_in_same_block(tmp_path: Path):
    sentence = "《广告法》第二十八条规定了虚假广告。"
    path = tmp_path / "repeated-anchor.docx"
    document = DocxDocument()
    document.add_paragraph(sentence + sentence)
    document.save(path)
    parsed = parse_docx(str(path))
    candidates = extract_rule_candidates(parsed, build_indexes(parsed), include_cases=False)
    claims = arbitrate_claim_candidates(candidates, parsed)
    assert [claim.source_locations[0].occurrence for claim in claims] == [0, 1]


def test_same_paragraph_inheritance_stops_after_intervening_sentence(tmp_path: Path):
    text = "《中华人民共和国商标法》第九条规定了有关规则。本案另有事实争议。第十条规定了其他规则。"
    path = tmp_path / "intervening-sentence.docx"
    document = DocxDocument()
    document.add_paragraph(text)
    document.save(path)
    parsed = parse_docx(str(path))
    candidates = extract_rule_candidates(parsed, build_indexes(parsed), include_cases=False)
    assert len(candidates) == 1
    assert candidates[0].entities.legal_sources[0].articles[0].article == "第九条"


def test_article_range_is_expanded_in_source_order():
    articles = extract_legal_sources("依据《民法典》第四十三条至第四十五条处理")[0].articles
    assert [article.article for article in articles] == [
        "第四十三条",
        "第四十四条",
        "第四十五条",
    ]


def test_bare_law_name_does_not_swallow_predicate_prefix():
    examples = {
        "人民法院可以认定为民法典第一百七十二条规定的相对人": "民法典",
        "原告请求依照反不正当竞争法第十七条处理": "反不正当竞争法",
        "该行为不属于反不正当竞争法第六条中的使用": "反不正当竞争法",
        "上述行为违反了反不正当竞争法第十二条规定": "反不正当竞争法",
    }
    for text, expected in examples.items():
        sources = extract_legal_sources(text)
        assert [source.title for source in sources] == [expected]


def test_case_patterns_are_bounded_and_support_guiding_case_without_di():
    refs = extract_case_refs("可参见指导案例262号及（2022）沪73知民初932号。")
    assert {ref.case_name for ref in refs if ref.case_name} == {"指导案例262号"}
    assert {ref.case_number for ref in refs if ref.case_number} == {"（2022）沪73知民初932号"}
    assert extract_case_refs("（2020）浙杭东证字第1712号公证书载明相关事实。") == []
    assert extract_case_refs("附录二：典型案例") == []


def test_multiple_clues_for_same_case_are_merged():
    guiding = extract_case_refs(
        "指导性案例262号：某科技有限公司诉某文化公司不正当竞争纠纷案"
    )
    assert len(guiding) == 1
    assert guiding[0].case_name.startswith("指导性案例262号：")

    numbered = extract_case_refs(
        "甲公司诉乙公司不正当竞争纠纷案（2023）浙民终1113号"
    )
    assert len(numbered) == 1
    assert numbered[0].case_number == "（2023）浙民终1113号"
    assert numbered[0].case_name == "甲公司诉乙公司不正当竞争纠纷案"


def test_claim_document_json_roundtrip_restores_typed_entities_and_context():
    doc = _parsed("依据《民法典》第五百七十七条，应核对引用。")
    candidates = extract_rule_candidates(doc, build_indexes(doc))
    claims = arbitrate_claim_candidates(candidates, doc)
    original = build_claim_document(doc, claims)
    restored = ClaimDocument.model_validate_json(original.model_dump_json())
    assert restored.claims[0].entities.legal_sources[0].title == "民法典"
    assert restored.claims[0].context_text == doc.anchors[0].text


class _EmptyCaseSearcher:
    def search_keyword(self, title: str, fulltext: str):
        return []

    def search_semantic(self, text: str):
        return []


def test_case_without_number_uses_both_routes_before_not_found():
    doc = _parsed("指导案例262号具有参考意义。")
    candidate = ClaimCandidate(
        claim_type=ClaimType.CASE_CITATION,
        anchor_ids=["line00001"],
        entities=CaseCitationEntities(
            case_refs=[
                CaseRef(
                    reference_type=CaseReferenceType.WITHOUT_CASE_NUMBER,
                    case_name="指导案例262号",
                )
            ]
        ),
    )
    claim_doc = build_claim_document(doc, arbitrate_claim_candidates([candidate], doc))
    checks = verify_case_claims(claim_doc, _EmptyCaseSearcher())
    assert len(checks) == 1
    assert checks[0].lookup_status == CaseLookupStatus.NOT_FOUND
    assert checks[0].cited_case_name == "指导案例262号"
    assert checks[0].source_attempts


def test_docx_footnote_is_parsed_into_stable_anchor(tmp_path: Path):
    path = tmp_path / "footnote.docx"
    document = DocxDocument()
    document.add_paragraph("正文。")
    document.save(path)
    footnotes_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
    <w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:footnote w:id="-1"><w:p><w:r><w:t>separator</w:t></w:r></w:p></w:footnote>
      <w:footnote w:id="2"><w:p><w:r><w:t>《民法典》第127条规定数据权益。</w:t></w:r></w:p></w:footnote>
    </w:footnotes>'''.encode("utf-8")
    rebuilt = tmp_path / "rebuilt.docx"
    with zipfile.ZipFile(path) as source, zipfile.ZipFile(rebuilt, "w") as target:
        for item in source.infolist():
            target.writestr(item, source.read(item.filename))
        target.writestr("word/footnotes.xml", footnotes_xml)
    parsed = parse_docx(str(rebuilt))
    note_blocks = [block for block in parsed.blocks if block.type == BlockType.FOOTNOTE]
    assert len(note_blocks) == 1
    assert note_blocks[0].note_id == "2"
    assert parsed.anchors[-1].note_type == "footnote"
    assert [anchor.anchor for anchor in parsed.anchors] == ["line00001", "line00002"]


def test_word_list_numbering_is_resolved(tmp_path: Path):
    path = tmp_path / "numbering.docx"
    document = DocxDocument()
    document.add_paragraph("第一项", style="List Number")
    document.add_paragraph("第二项", style="List Number")
    document.save(path)
    parsed = parse_docx(str(path))
    assert [block.numbering_text for block in parsed.blocks] == ["1.", "2."]
    assert all(not block.numbering_unresolved for block in parsed.blocks)


def test_merged_table_cell_is_not_duplicated(tmp_path: Path):
    path = tmp_path / "merged-table.docx"
    document = DocxDocument()
    table = document.add_table(rows=1, cols=2)
    merged = table.cell(0, 0).merge(table.cell(0, 1))
    merged.text = "《民法典》第127条"
    document.save(path)
    parsed = parse_docx(str(path))
    table_blocks = [block for block in parsed.blocks if block.type == BlockType.TABLE_CELL]
    assert len(table_blocks) == 1
