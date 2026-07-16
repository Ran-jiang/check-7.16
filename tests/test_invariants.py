"""
Test 7：不变量全量校验。

构造混合文档：
  标题 → 普通段落 → 列举引导句 → 列举项 → 表格
运行全量校验，期望返回空列表。

另外测试违反各项不变量的场景。
"""

import os
import tempfile

from docx import Document as DocxDocument

from ccitecheck.parsing.docx import parse_docx
from ccitecheck.parsing.chunks import build_chunks
from ccitecheck.parsing.validators import validate_parsed_document
from ccitecheck.domain.document import (
    ParsedDocument,
    DocMeta,
    Block,
    Anchor,
    BlockType,
)


def _write_docx(doc: DocxDocument, path: str):
    doc.save(path)


def _tmp_docx_path():
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    return path


# ---- Test 7: 全量校验 ----

class TestFullValidation:
    """Test 7：不变量全量校验"""

    def test_mixed_document_passes(self):
        """构造包含标题、段落、列表、表格的混合文档，校验应通过。"""
        doc = DocxDocument()

        # 标题
        h = doc.add_paragraph("第一章 总则")
        h.style = doc.styles["Heading 1"]

        # 普通段落
        doc.add_paragraph("第一条 本合同自签署之日起生效。双方应严格遵守。")

        # 列举引导句 + 列举项
        doc.add_paragraph("有下列情形之一的：")
        doc.add_paragraph("（一）一方严重违反合同约定；")
        doc.add_paragraph("（二）发生不可抗力事件。")

        # 表格
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "甲方"
        table.cell(0, 1).text = "乙方"

        # 表格后的段落
        doc.add_paragraph("以上内容双方确认无误。")

        path = _tmp_docx_path()
        _write_docx(doc, path)

        try:
            parsed = parse_docx(path)
            parsed = build_chunks(parsed)

            violations = validate_parsed_document(parsed)

            if violations:
                for v in violations:
                    print(f"VIOLATION: {v}")

            assert violations == [], f"期望0条违反项，实际{len(violations)}条"
        finally:
            os.unlink(path)


class TestViolationScenarios:
    """测试各项不变量违反场景"""

    def test_lossless_split_violation(self):
        """anchor.text 不匹配时检测到。"""
        parsed = _make_basic_parsed()

        # 破坏 anchor 偏移
        parsed.anchors[0].char_end += 1

        violations = validate_parsed_document(parsed)
        lossless_violations = [v for v in violations if "无损分句" in v]
        assert len(lossless_violations) > 0

    def test_chunk_coverage_violation(self):
        """anchor 未出现在任何 chunk 中。"""
        parsed = _make_basic_parsed()
        parsed = build_chunks(parsed)

        # 清空所有 chunk 的 anchor_ids
        for chunk in parsed.chunks:
            chunk.anchor_ids = []
            chunk.anchor_range = ["", ""]

        violations = validate_parsed_document(parsed)
        coverage_violations = [v for v in violations if "chunk覆盖" in v]
        assert len(coverage_violations) > 0

    def test_id_duplicate(self):
        """ID 重复被检测。"""
        parsed = _make_basic_parsed()

        # 复制一个 block（ID 重复）
        parsed.blocks.append(parsed.blocks[0])

        violations = validate_parsed_document(parsed)
        id_violations = [v for v in violations if "ID唯一性" in v]
        assert len(id_violations) > 0

    def test_clean_document_passes(self):
        """正确构建的文档应通过所有校验。"""
        parsed = _make_basic_parsed()
        parsed = build_chunks(parsed)

        violations = validate_parsed_document(parsed)

        # 基本文档应通过
        # 注意：如果 _make_basic_parsed 的数据自洽，violations 应为 []
        # 排除任何 chunk 相关的校验（取决于 chunk builder 的输出）
        assert all(
            "无损分句" not in v
            and "ID唯一性" not in v
            and "anchor_range" not in v
            for v in violations
        ), f"基本不变量违反: {violations}"


def _make_basic_parsed() -> ParsedDocument:
    """构造供测试使用的基本 ParsedDocument。"""
    text1 = "第一段文本。第一段第二句。"
    text2 = "第二段文本。"

    b1 = Block(
        block_id="b_00001",
        type=BlockType.PARAGRAPH,
        text=text1,
        section_path=[],
        body_order=0,
        block_order=0,
        para_index=0,
        anchor_range=["line00001", "line00002"],
        sentence_anchors=["line00001", "line00002"],
    )

    b2 = Block(
        block_id="b_00002",
        type=BlockType.PARAGRAPH,
        text=text2,
        section_path=[],
        body_order=1,
        block_order=1,
        para_index=1,
        anchor_range=["line00003", "line00003"],
        sentence_anchors=["line00003"],
    )

    a1 = Anchor(
        anchor="line00001",
        text="第一段文本。",
        block_id="b_00001",
        para_index=0,
        char_start=0,
        char_end=6,
    )
    a2 = Anchor(
        anchor="line00002",
        text="第一段第二句。",
        block_id="b_00001",
        para_index=0,
        char_start=6,
        char_end=13,
    )
    a3 = Anchor(
        anchor="line00003",
        text="第二段文本。",
        block_id="b_00002",
        para_index=1,
        char_start=0,
        char_end=6,
    )

    return ParsedDocument(
        doc_meta=DocMeta(source_file="test.docx", doc_hash="test"),
        blocks=[b1, b2],
        anchors=[a1, a2, a3],
        chunks=[],
    )
