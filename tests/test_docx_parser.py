"""
测试 DOCX 解析器。

使用 python-docx 在测试内动态构造 DOCX fixture。

测试覆盖：
  - 基本段落解析
  - 表格顺序（Test 4）
  - 空段落过滤
  - 标题识别
  - 列举项
  - body_order / block_order / para_index 正确性
  - 软换行和制表符归一化
"""

import os
import tempfile

from docx import Document as DocxDocument

from parser.docx_parser import (
    parse_docx,
)
from parser.schema import BlockType
from parser.utils import is_empty_text


# ---- 辅助函数 ----

def _write_docx(doc: DocxDocument, path: str):
    """将 python-docx Document 写入临时文件。"""
    doc.save(path)


def _tmp_docx_path():
    """生成临时 DOCX 文件路径。"""
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    return path


# ---- Test 4: 表格顺序 ----

class TestTableOrder:
    """Test 4：表格顺序"""

    def test_table_between_paragraphs(self):
        """段落1, 表格, 段落2 → 阅读顺序为 段落1, 表格单元格, 段落2"""
        doc = DocxDocument()
        doc.add_paragraph("段落1")
        # 添加表格
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "表格单元格内容"
        doc.add_paragraph("段落2")

        path = _tmp_docx_path()
        _write_docx(doc, path)

        try:
            parsed = parse_docx(path)

            # 应有 3 个 block
            assert len(parsed.blocks) == 3

            # 顺序检查
            b0 = parsed.blocks[0]
            b1 = parsed.blocks[1]
            b2 = parsed.blocks[2]

            assert b0.text == "段落1"
            assert b0.type == BlockType.PARAGRAPH
            assert b0.body_order == 0
            assert b0.block_order == 0

            assert b1.text == "表格单元格内容"
            assert b1.type == BlockType.TABLE_CELL
            assert b1.table_index == 0
            assert b1.row_index == 0
            assert b1.cell_index == 0
            # 表格 cell 的 body_order 等于表格的顶层 body order (1)
            assert b1.body_order == 1
            assert b1.block_order == 1
            assert b1.para_index is None  # table_cell 无 para_index

            assert b2.text == "段落2"
            assert b2.type == BlockType.PARAGRAPH
            assert b2.body_order == 2
            assert b2.block_order == 2
        finally:
            os.unlink(path)

    def test_multiple_cells_in_table(self):
        """同一表格多个非空 cell，共享 body_order。"""
        doc = DocxDocument()
        doc.add_paragraph("段落1")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "R0C0"
        table.cell(0, 1).text = "R0C1"
        table.cell(1, 0).text = "R1C0"
        # R1C1 为空
        table.cell(1, 1).text = ""
        doc.add_paragraph("段落2")

        path = _tmp_docx_path()
        _write_docx(doc, path)

        try:
            parsed = parse_docx(path)

            # 3 个非空 cell + 2 个段落 = 5 blocks
            assert len(parsed.blocks) == 5

            # 表格 cells 的 body_order 应相同（都是表格的顶层顺序 = 1）
            cell_blocks = [b for b in parsed.blocks if b.type == BlockType.TABLE_CELL]
            assert len(cell_blocks) == 3
            body_orders = {b.body_order for b in cell_blocks}
            assert len(body_orders) == 1  # 所有 cell 共享同一 body_order
            assert 1 in body_orders

            # block_order 应互不相同且单调递增
            all_orders = [b.block_order for b in parsed.blocks]
            assert all_orders == sorted(all_orders)
            assert len(set(all_orders)) == len(all_orders)
        finally:
            os.unlink(path)


class TestParsing:
    """基础解析测试"""

    def test_paragraph_parsing(self):
        """普通段落解析。"""
        doc = DocxDocument()
        doc.add_paragraph("第一段文本。第二句文本。")

        path = _tmp_docx_path()
        _write_docx(doc, path)

        try:
            parsed = parse_docx(path)

            assert len(parsed.blocks) == 1
            b = parsed.blocks[0]
            assert b.type == BlockType.PARAGRAPH
            assert b.text == "第一段文本。第二句文本。"
            assert b.para_index == 0
            assert b.body_order == 0
            assert b.block_order == 0

            # 应有 2 个 anchors
            assert len(parsed.anchors) == 2
            assert parsed.anchors[0].anchor == "line00001"
            assert parsed.anchors[1].anchor == "line00002"
        finally:
            os.unlink(path)

    def test_empty_paragraph_filtered(self):
        """空段落被过滤但不影响 para_index。"""
        doc = DocxDocument()
        doc.add_paragraph("段落1")
        doc.add_paragraph("")  # 空段落
        doc.add_paragraph("段落2")

        path = _tmp_docx_path()
        _write_docx(doc, path)

        try:
            parsed = parse_docx(path)

            # 2 个非空 blocks
            assert len(parsed.blocks) == 2

            # para_index 不重排：空段落占用了 para_index=1
            assert parsed.blocks[0].para_index == 0
            assert parsed.blocks[1].para_index == 2
        finally:
            os.unlink(path)

    def test_heading_detection(self):
        """标题识别与 section_path。"""
        doc = DocxDocument()
        # 使用标题样式
        h1 = doc.add_paragraph("第一章 总则")
        h1.style = doc.styles["Heading 1"]
        doc.add_paragraph("第一条 内容文本。")

        path = _tmp_docx_path()
        _write_docx(doc, path)

        try:
            parsed = parse_docx(path)

            assert len(parsed.blocks) == 2
            h_block = parsed.blocks[0]
            p_block = parsed.blocks[1]

            assert h_block.type == BlockType.HEADING
            assert h_block.section_path == ["第一章 总则"]

            # 后续段落的 section_path 继承标题
            assert p_block.section_path == ["第一章 总则"]
            assert p_block.type == BlockType.PARAGRAPH
        finally:
            os.unlink(path)

    def test_article_start_detection(self):
        """'第X条' is_article_start 检测。"""
        doc = DocxDocument()
        doc.add_paragraph("第三十七条 劳动者提前三十日以书面形式通知用人单位。")

        path = _tmp_docx_path()
        _write_docx(doc, path)

        try:
            parsed = parse_docx(path)

            assert len(parsed.blocks) == 1
            b = parsed.blocks[0]
            assert b.is_article_start is True
            assert b.type == BlockType.PARAGRAPH  # 不是 heading
        finally:
            os.unlink(path)

    def test_list_item_detection(self):
        """列举项检测。"""
        doc = DocxDocument()
        doc.add_paragraph("有下列情形之一的：")
        doc.add_paragraph("（一）劳动者严重违反规章制度；")
        doc.add_paragraph("（二）劳动者严重失职；")

        path = _tmp_docx_path()
        _write_docx(doc, path)

        try:
            parsed = parse_docx(path)

            assert len(parsed.blocks) == 3
            b0 = parsed.blocks[0]
            b1 = parsed.blocks[1]
            b2 = parsed.blocks[2]

            assert b1.is_list_item is True
            assert b1.type == BlockType.LIST_ITEM
            assert b2.is_list_item is True

            # 三个 block 应共享同一 list_group_id
            assert b0.list_group_id is not None
            assert b0.list_group_id == b1.list_group_id
            assert b1.list_group_id == b2.list_group_id
        finally:
            os.unlink(path)


class TestNormalization:
    """文本归一化测试"""

    def test_whitespace_normalization(self):
        """空白字符归一化。"""
        from parser.utils import normalize_whitespace as norm

        assert norm("hello  world") == "hello world"
        assert norm("  leading") == "leading"
        assert norm("trailing  ") == "trailing"
        assert norm("") == ""

    def test_is_empty_text(self):
        """空白文本判断。"""
        assert is_empty_text("") is True
        assert is_empty_text("  ") is True
        assert is_empty_text("\t\n") is True
        assert is_empty_text("你好") is False


class TestNumbering:
    """自动编号测试"""

    def test_has_numbering_detection(self):
        """检测 has_numbering。"""
        doc = DocxDocument()
        doc.add_paragraph("普通段落")

        path = _tmp_docx_path()
        _write_docx(doc, path)

        try:
            parsed = parse_docx(path)
            b = parsed.blocks[0]
            # 无编号属性的段落
            assert b.has_numbering is False
        finally:
            os.unlink(path)
