from datetime import datetime, timedelta, timezone
from io import BytesIO

from docx import Document

from apps.api.web_sessions import WebDocumentSession, render_revised_docx
from ccitecheck.domain.revisions import RevisionProposal


def test_render_revised_docx_applies_only_accepted_safe_revision():
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("依据《民法典》")
    paragraph.add_run("第九款处理。")
    source = BytesIO()
    document.save(source)
    proposal = RevisionProposal(
        strategy="replace_exact_text",
        original_text="依据《民法典》第九款处理。",
        revised_text="依据《民法典》第三款处理。",
        rationale="款号定位错误",
        machine_applicable=True,
        preconditions=["original_text_unique"],
    )
    session = WebDocumentSession(
        session_id="session",
        file_name="test.docx",
        document_bytes=source.getvalue(),
        revisions={"sc_00001": proposal},
        expires_at=datetime.now(timezone.utc) + timedelta(hours=1),
        accepted={"sc_00001"},
    )

    revised = Document(BytesIO(render_revised_docx(session)))

    assert revised.paragraphs[0].text == "依据《民法典》第三款处理。"


def test_render_revised_docx_rejects_ambiguous_original_text():
    document = Document()
    document.add_paragraph("第九款。")
    document.add_paragraph("第九款。")
    source = BytesIO()
    document.save(source)
    proposal = RevisionProposal(
        strategy="replace_exact_text",
        original_text="第九款。",
        revised_text="第三款。",
        rationale="款号定位错误",
        machine_applicable=True,
    )
    session = WebDocumentSession(
        session_id="session",
        file_name="test.docx",
        document_bytes=source.getvalue(),
        revisions={"sc_00001": proposal},
        expires_at=datetime.now(timezone.utc) + timedelta(hours=1),
        accepted={"sc_00001"},
    )

    try:
        render_revised_docx(session)
    except ValueError as error:
        assert "无法唯一定位" in str(error)
    else:
        raise AssertionError("ambiguous revision must fail")
