"""
测试 Chunk 构建器。

测试覆盖：
  - 短 block 合并（Test 5）
  - chunk 不重复存文本
  - chunk 连续性
  - 超长 block 拆分
"""

from docx import Document as DocxDocument

from parser.docx_parser import parse_docx
from parser.chunk_builder import build_chunks

import os
import tempfile


def _write_docx(doc: DocxDocument, path: str):
    doc.save(path)


def _tmp_docx_path():
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    return path


# ---- Test 5: chunk 不重复 ----

class TestChunkNoOverlap:
    """Test 5：chunk 不重复"""

    def test_five_sentences_one_chunk(self):
        """5 个短句的段落生成一个 chunk，无重叠。"""
        doc = DocxDocument()
        text = "第一句。第二句。第三句。第四句。第五句。"
        doc.add_paragraph(text)

        path = _tmp_docx_path()
        _write_docx(doc, path)

        try:
            parsed = parse_docx(path)
            parsed = build_chunks(parsed)

            # 5 个短句，总共 token 数 < 1500，应生成 1 个 chunk
            assert len(parsed.chunks) == 1

            chunk = parsed.chunks[0]
            # 所有 5 个 anchor 都在这个 chunk 中
            assert len(chunk.anchor_ids) == 5
            # 无重叠
            assert chunk.overlap_anchor_ids == []
        finally:
            os.unlink(path)

    def test_no_sliding_window(self):
        """不使用滑动窗口，每个 anchor 只属于一个 chunk。"""
        doc = DocxDocument()
        # 5 个短句
        text = "第一句。第二句。第三句。第四句。第五句。"
        doc.add_paragraph(text)

        path = _tmp_docx_path()
        _write_docx(doc, path)

        try:
            parsed = parse_docx(path)
            parsed = build_chunks(parsed)

            # 收集所有 anchor 出现的 chunk 数
            anchor_count: dict[str, int] = {}
            for chunk in parsed.chunks:
                for aid in chunk.anchor_ids:
                    anchor_count[aid] = anchor_count.get(aid, 0) + 1

            # 每个 anchor 应只出现在一个 chunk 中
            # （5个短句不会触发超长block拆分）
            for aid, count in anchor_count.items():
                assert count == 1, f"anchor {aid} 出现在 {count} 个 chunk 中"
        finally:
            os.unlink(path)


class TestChunkBasics:
    """Chunk 基本规则测试"""

    def test_chunk_with_section_path(self):
        """chunk 继承 section_path。"""
        doc = DocxDocument()
        h = doc.add_paragraph("第一章 总则")
        h.style = doc.styles["Heading 1"]
        doc.add_paragraph("第一条 内容一。内容二。")

        path = _tmp_docx_path()
        _write_docx(doc, path)

        try:
            parsed = parse_docx(path)
            parsed = build_chunks(parsed)

            assert len(parsed.chunks) >= 1

            # chunk 应包含 section_path
            chunk = parsed.chunks[0]
            assert "第一章 总则" in chunk.section_path
        finally:
            os.unlink(path)

    def test_empty_document(self):
        """空文档不报错。"""
        doc = DocxDocument()
        path = _tmp_docx_path()
        _write_docx(doc, path)

        try:
            parsed = parse_docx(path)
            parsed = build_chunks(parsed)

            assert parsed.chunks == []
        finally:
            os.unlink(path)


class TestLongBlockSplit:
    """超长 block 拆分测试"""

    def test_long_block_split_with_overlap(self):
        """超长 block 按 sentence anchors 拆分，有重叠。"""
        doc = DocxDocument()
        # 构造一个超长段落：很多句子
        sentences = []
        for i in range(200):
            sentences.append(f"第{i}条内容文本描述与法律规范。")
        text = "".join(sentences)
        doc.add_paragraph(text)

        path = _tmp_docx_path()
        _write_docx(doc, path)

        try:
            parsed = parse_docx(path)
            parsed = build_chunks(parsed)

            # 应有多个 chunks
            assert len(parsed.chunks) > 1

            # 至少有一个 chunk 有 overlap_anchor_ids（超长block拆分）
            # 注意：如果 block 没有超过 chunk 上限，就不会拆分
            # 这个测试中 200 句中文字符肯定会超过上限
            has_overlap = any(
                len(c.overlap_anchor_ids) > 0 for c in parsed.chunks
            )
            assert has_overlap, "超长 block 应产生重叠 chunk"
        finally:
            os.unlink(path)
