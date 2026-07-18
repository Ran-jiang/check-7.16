import base64
import importlib
from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from ccitecheck.infrastructure.database import connect, init_db, upsert_article, upsert_law


def _reject_named_temporary_file(*args, **kwargs):
    raise AssertionError("API must use a closed, reopenable temporary DOCX path")


def test_word_addin_document_check_api(tmp_path, monkeypatch):
    db_path = tmp_path / "laws.sqlite"
    init_db(db_path)
    with connect(db_path) as connection:
        law_id = upsert_law(
            connection,
            {"title": "中华人民共和国民法典", "source_type": "law"},
        )
        upsert_article(
            connection,
            law_id,
            {
                "article_no": "第五百七十七条",
                "text": "当事人一方不履行合同义务或者履行合同义务不符合约定的，应当承担违约责任。",
            },
        )

    document = Document()
    document.add_paragraph("依据《中华人民共和国民法典》第五百七十七条，被告应当承担违约责任。")
    buffer = BytesIO()
    document.save(buffer)

    api_module = importlib.import_module("apps.api.app")
    monkeypatch.setattr(api_module, "LAW_DB", db_path)
    monkeypatch.setattr(api_module.tempfile, "NamedTemporaryFile", _reject_named_temporary_file)
    client = TestClient(api_module.app)
    response = client.post(
        "/api/checks",
        json={
            "file_name": "test.docx",
            "docx_base64": base64.b64encode(buffer.getvalue()).decode(),
            "semantic_check": False,
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert "legal_checks" not in payload["verification"]
    assert payload["summary"]["card_total"] == 1
    assert payload["summary"]["reference_total"] == 1
    assert payload["semantic_check"] is False
    assert payload["summary"]["total"] == 1


def test_web_document_check_returns_preview_session_and_download(tmp_path, monkeypatch):
    db_path = tmp_path / "laws.sqlite"
    _seed_law_db(db_path)
    document = Document()
    document.add_paragraph("依据《中华人民共和国民法典》第五百七十七条，被告应当承担违约责任。")
    buffer = BytesIO()
    document.save(buffer)

    api_module = importlib.import_module("apps.api.app")
    monkeypatch.setattr(api_module, "LAW_DB", db_path)
    monkeypatch.setattr(api_module, "create_run", lambda *args, **kwargs: (_ for _ in ()).throw(AssertionError("web uploads must not be debug-captured")))
    client = TestClient(api_module.app)
    response = client.post("/api/web/checks", json={
        "file_name": "网页测试.docx",
        "docx_base64": base64.b64encode(buffer.getvalue()).decode(),
        "semantic_check": False,
    })

    assert response.status_code == 200
    payload = response.json()
    assert payload["session_id"]
    assert payload["preview_blocks"][0]["block_id"] == "word:p:0"
    assert payload["preview_blocks"][0]["text"].startswith("依据《中华人民共和国民法典》")
    download = client.get(f"/api/web/sessions/{payload['session_id']}/document")
    assert download.status_code == 200
    assert download.content.startswith(b"PK")


def test_web_text_check_uses_same_pipeline(tmp_path, monkeypatch):
    db_path = tmp_path / "laws.sqlite"
    _seed_law_db(db_path)
    api_module = importlib.import_module("apps.api.app")
    monkeypatch.setattr(api_module, "LAW_DB", db_path)
    client = TestClient(api_module.app)

    response = client.post("/api/web/checks/text", json={
        "text": "依据《中华人民共和国民法典》第五百七十七条，被告应当承担违约责任。",
        "semantic_check": False,
    })

    assert response.status_code == 200
    assert response.json()["summary"]["total"] == 1
    assert response.json()["preview_blocks"][0]["block_id"] == "word:p:0"


def test_web_page_has_public_security_and_cache_headers():
    api_module = importlib.import_module("apps.api.app")
    client = TestClient(api_module.app)

    response = client.get("/web/")

    assert response.status_code == 200
    assert response.headers["cache-control"] == "no-cache"
    assert "frame-ancestors 'none'" in response.headers["content-security-policy"]
    assert response.headers["x-content-type-options"] == "nosniff"


def _seed_law_db(db_path):
    init_db(db_path)
    with connect(db_path) as connection:
        law_id = upsert_law(
            connection,
            {"title": "中华人民共和国民法典", "source_type": "law"},
        )
        upsert_article(
            connection,
            law_id,
            {
                "article_no": "第五百七十七条",
                "text": "当事人一方不履行合同义务或者履行合同义务不符合约定的，应当承担违约责任。",
            },
        )


def test_selection_check_api(tmp_path, monkeypatch):
    db_path = tmp_path / "laws.sqlite"
    _seed_law_db(db_path)

    api_module = importlib.import_module("apps.api.app")
    monkeypatch.setattr(api_module, "LAW_DB", db_path)
    monkeypatch.setattr(api_module.tempfile, "NamedTemporaryFile", _reject_named_temporary_file)
    client = TestClient(api_module.app)
    response = client.post(
        "/api/checks/selection",
        json={
            "file_name": "test.docx",
            "text": "依据《中华人民共和国民法典》第五百七十七条，被告应当承担违约责任。",
            "source_blocks": [{"block_id": "word:p:7", "char_start": 12}],
            "semantic_check": False,
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["file_name"] == "test.docx（选中片段）"
    assert payload["summary"]["total"] == 1
    assert payload["verification"]["statute_results"][0]["lookup_status"] == "article_found"
    location = payload["verification"]["statute_results"][0]["source_locations"][0]
    assert location["block_id"] == "word:p:7"
    assert location["char_start"] == 12


def test_selection_check_rejects_empty_text(tmp_path, monkeypatch):
    api_module = importlib.import_module("apps.api.app")
    client = TestClient(api_module.app)
    response = client.post(
        "/api/checks/selection",
        json={"file_name": "test.docx", "text": "   \n  ", "semantic_check": False},
    )
    assert response.status_code == 400


def test_case_only_selection_reports_unconfigured_case_source(tmp_path, monkeypatch):
    api_module = importlib.import_module("apps.api.app")
    verification_module = importlib.import_module("ccitecheck.application.verify_claims")
    from ccitecheck.tracing.sources.pkulaw.client import PkulawNotConfiguredError

    class UnconfiguredCaseSource:
        def search_keyword(self, title, fulltext):
            raise PkulawNotConfiguredError("案例数据源未配置")

        def search_semantic(self, text):
            raise PkulawNotConfiguredError("案例数据源未配置")

    monkeypatch.setattr(verification_module, "PkulawCaseSource", UnconfiguredCaseSource)
    monkeypatch.setattr(api_module, "LAW_DB", tmp_path / "laws.sqlite")
    client = TestClient(api_module.app)
    response = client.post(
        "/api/checks/selection",
        json={
            "file_name": "案例研究.docx",
            "text": "指导案例262号具有参考意义。",
            "semantic_check": True,
            "include_statutes": False,
            "include_cases": True,
        },
    )
    assert response.status_code == 200
    payload = response.json()
    assert payload["summary"]["total"] == 1
    assert payload["summary"]["bugs"] == 1
    assert payload["verification"]["statute_results"] == []
    assert payload["verification"]["case_results"][0]["lookup_status"] == "source_not_configured"
    assert payload["document_key"].startswith("sha256:")


def test_scope_validation_requires_at_least_one(tmp_path, monkeypatch):
    api_module = importlib.import_module("apps.api.app")
    client = TestClient(api_module.app)
    response = client.post(
        "/api/checks/selection",
        json={
            "file_name": "t.docx",
            "text": "依据《中华人民共和国民法典》第五百七十七条。",
            "semantic_check": False,
            "include_statutes": False,
            "include_cases": False,
        },
    )
    assert response.status_code == 400


def test_statutes_can_be_excluded(tmp_path, monkeypatch):
    db_path = tmp_path / "laws.sqlite"
    _seed_law_db(db_path)
    api_module = importlib.import_module("apps.api.app")
    monkeypatch.setattr(api_module, "LAW_DB", db_path)
    client = TestClient(api_module.app)
    response = client.post(
        "/api/checks/selection",
        json={
            "file_name": "t.docx",
            "text": "依据《中华人民共和国民法典》第五百七十七条，应当承担违约责任。",
            "semantic_check": False,
            "include_statutes": False,
            "include_cases": False or True,
        },
    )
    assert response.status_code == 200
    payload = response.json()
    assert payload["verification"]["statute_results"] == []
