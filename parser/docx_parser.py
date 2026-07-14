"""
CCitecheck v0.1 DOCX 解析器。

使用 python-docx 解析 DOCX 文件，生成三层结构化数据。

核心逻辑：
  1. 遍历 DOCX body 的 XML 子元素（w:p 和 w:tbl），保持原始顺序
  2. 软换行 w:br 替换为半角空格，制表符替换为半角空格
  3. 过滤空段落和空表格单元格
  4. 表格 cell 各自成为一个 table_cell block
  5. 分句、标题识别、编号检测
"""

from __future__ import annotations

import hashlib
import zipfile
from typing import Optional

from docx import Document as DocxDocument
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.table import Table as DocxTable

from .heading_detector import (
    detect_heading,
    scan_chapter_types,
)
from .schema import (
    Anchor,
    Block,
    BlockType,
    DocMeta,
    HeadingSource,
    ParsedDocument,
)
from .sentence_splitter import split_sentences
from .utils import (
    detect_chinese_list_item,
    is_article_start,
    is_empty_text,
    make_id_counter,
    normalize_whitespace,
)


# ---- 常量 ----
# 软换行 XML 标签
TAG_BR = qn("w:br")
# 制表符 XML 标签
TAG_TAB = qn("w:tab")
# 段落 XML 标签
TAG_P = qn("w:p")
# 表格 XML 标签
TAG_TBL = qn("w:tbl")


# ---- 段落文本提取 ----
# 需要从 w:p 的 w:r 子元素中提取文本，并将 w:br 替换为空格、w:tab 替换为空格


def extract_paragraph_text(para_element) -> str:
    """
    从 w:p XML 元素中提取文本。

    规则：
      - w:br 替换为一个半角空格
      - w:tab 替换为一个半角空格
      - 多个 w:r 的文本拼接

    Args:
        para_element: lxml 元素（w:p）

    Returns:
        提取后的文本（尚未归一化）
    """
    parts: list[str] = []
    # 遍历所有子元素
    for child in para_element.iter():
        if child.tag == TAG_BR:
            # 软换行 → 半角空格
            parts.append(" ")
        elif child.tag == TAG_TAB:
            # 制表符 → 半角空格
            parts.append(" ")
        elif child.tag == qn("w:t"):
            # 文本内容
            if child.text:
                parts.append(child.text)
    return "".join(parts)


def extract_cell_text(cell) -> str:
    """
    从表格单元格中提取文本。

    cell 内可能包含多个非空段落，用一个半角空格拼接。
    每个段落内的软换行和制表符同样归一化。

    Args:
        cell: python-docx Cell 对象

    Returns:
        提取并归一化后的文本
    """
    para_texts: list[str] = []
    for para in cell.paragraphs:
        # 通过 para._element 获取底层 XML 以正确处理软换行
        raw_text = extract_paragraph_text(para._element)
        text = normalize_whitespace(raw_text)
        if text:
            para_texts.append(text)
    return " ".join(para_texts)


# ---- 自动编号处理 ----
# 检测段落是否携带编号属性 w:numPr


class NumberingResolver:
    """按 numbering.xml 的定义还原常见 Word 自动编号。"""

    def __init__(self, docx: DocxDocument):
        self.levels: dict[tuple[str, int], dict[str, object]] = {}
        self.counters: dict[str, dict[int, int]] = {}
        try:
            root = docx.part.numbering_part.element
        except (AttributeError, NotImplementedError):
            return
        abstract_by_id = {
            node.get(qn("w:abstractNumId")): node
            for node in root.findall(qn("w:abstractNum"))
        }
        for num in root.findall(qn("w:num")):
            num_id = num.get(qn("w:numId"))
            abstract_ref = num.find(qn("w:abstractNumId"))
            if num_id is None or abstract_ref is None:
                continue
            abstract = abstract_by_id.get(abstract_ref.get(qn("w:val")))
            if abstract is None:
                continue
            for level in abstract.findall(qn("w:lvl")):
                ilvl = int(level.get(qn("w:ilvl"), "0"))
                start = level.find(qn("w:start"))
                num_fmt = level.find(qn("w:numFmt"))
                lvl_text = level.find(qn("w:lvlText"))
                self.levels[(num_id, ilvl)] = {
                    "start": int(start.get(qn("w:val"), "1")) if start is not None else 1,
                    "format": num_fmt.get(qn("w:val"), "decimal") if num_fmt is not None else "decimal",
                    "text": lvl_text.get(qn("w:val"), f"%{ilvl + 1}.") if lvl_text is not None else f"%{ilvl + 1}.",
                }

    def resolve(self, num_pr) -> tuple[bool, Optional[str], bool]:
        if num_pr is None:
            return False, None, False
        num_id_element = num_pr.find(qn("w:numId"))
        level_element = num_pr.find(qn("w:ilvl"))
        if num_id_element is None:
            return True, None, True
        num_id = num_id_element.get(qn("w:val"))
        if not num_id or num_id == "0":
            return False, None, False
        ilvl = int(level_element.get(qn("w:val"), "0")) if level_element is not None else 0
        definition = self.levels.get((num_id, ilvl))
        if definition is None:
            return True, None, True

        counters = self.counters.setdefault(num_id, {})
        for deeper in [level for level in counters if level > ilvl]:
            del counters[deeper]
        current = counters.get(ilvl, int(definition["start"]) - 1) + 1
        counters[ilvl] = current
        rendered = str(definition["text"])
        for level in range(9):
            placeholder = f"%{level + 1}"
            if placeholder not in rendered:
                continue
            level_definition = self.levels.get((num_id, level))
            value = counters.get(level)
            if level_definition is None or value is None:
                return True, None, True
            formatted = _format_number(value, str(level_definition["format"]))
            if formatted is None:
                return True, None, True
            rendered = rendered.replace(placeholder, formatted)
        return True, rendered, False


def _format_number(value: int, number_format: str) -> Optional[str]:
    if number_format in {"decimal", "decimalZero"}:
        return str(value)
    if number_format in {"lowerLetter", "upperLetter"} and 1 <= value <= 26:
        letter = chr(ord("a") + value - 1)
        return letter.upper() if number_format == "upperLetter" else letter
    if number_format in {"lowerRoman", "upperRoman"} and 1 <= value <= 3999:
        pairs = ((1000, "M"), (900, "CM"), (500, "D"), (400, "CD"), (100, "C"),
                 (90, "XC"), (50, "L"), (40, "XL"), (10, "X"), (9, "IX"),
                 (5, "V"), (4, "IV"), (1, "I"))
        rest, result = value, ""
        for amount, glyph in pairs:
            while rest >= amount:
                result += glyph
                rest -= amount
        return result if number_format == "upperRoman" else result.lower()
    if number_format in {"chineseCounting", "chineseCountingThousand", "ideographTraditional"}:
        return _int_to_chinese(value)
    if number_format == "bullet":
        return "•"
    return None


def _int_to_chinese(value: int) -> str:
    if not 0 < value < 10000:
        return str(value)
    digits = "零一二三四五六七八九"
    units = ((1000, "千"), (100, "百"), (10, "十"))
    result, rest, zero_pending = "", value, False
    for amount, label in units:
        digit, rest = divmod(rest, amount)
        if digit:
            if zero_pending and result:
                result += "零"
            if not (amount == 10 and digit == 1 and not result):
                result += digits[digit]
            result += label
            zero_pending = False
        elif result and rest:
            zero_pending = True
    if rest:
        if zero_pending:
            result += "零"
        result += digits[rest]
    return result


def _paragraph_num_pr(para_element, docx: DocxDocument):
    p_pr = para_element.find(qn("w:pPr"))
    direct = p_pr.find(qn("w:numPr")) if p_pr is not None else None
    if direct is not None:
        return direct
    style_name = _get_style_name(para_element, docx)
    if not style_name:
        return None
    try:
        style = docx.styles[style_name]
        style_p_pr = style._element.find(qn("w:pPr"))
        return style_p_pr.find(qn("w:numPr")) if style_p_pr is not None else None
    except KeyError:
        return None


def check_paragraph_numbering(
    para_element,
    docx: DocxDocument | None = None,
    resolver: NumberingResolver | None = None,
) -> tuple[bool, Optional[str], bool]:
    """
    检测段落是否携带自动编号。

    Args:
        para_element: lxml 元素（w:p）

    Returns:
        (has_numbering, numbering_text, numbering_unresolved)
    """
    if docx is None or resolver is None:
        p_pr = para_element.find(qn("w:pPr"))
        num_pr = p_pr.find(qn("w:numPr")) if p_pr is not None else None
        return (num_pr is not None, None, num_pr is not None)
    return resolver.resolve(_paragraph_num_pr(para_element, docx))


# ---- 主解析函数 ----


def parse_docx(file_path: str) -> ParsedDocument:
    """
    解析 DOCX 文件，生成 ParsedDocument。

    这是整个解析流程的入口函数。

    Args:
        file_path: DOCX 文件路径

    Returns:
        ParsedDocument 对象

    Raises:
        FileNotFoundError: 文件不存在
        ValueError: 文件不是有效 DOCX
    """
    # ---- 读取文件并计算 SHA-256 ----
    with open(file_path, "rb") as f:
        raw_bytes = f.read()
    doc_hash = "sha256:" + hashlib.sha256(raw_bytes).hexdigest()

    # ---- 使用 python-docx 打开 ----
    docx = DocxDocument(file_path)
    numbering_resolver = NumberingResolver(docx)

    # ---- ID 计数器 ----
    next_anchor_id = make_id_counter("line", 5)
    next_block_id = make_id_counter("b_", 5)
    next_list_group_id = make_id_counter("lg_", 5)

    # ---- 第一遍扫描：收集所有段落文本用于章节类型扫描 ----
    # 遍历 body 元素收集所有 w:p 文本
    body = docx.element.body
    all_para_texts: list[str] = []
    for child in body:
        if child.tag == TAG_P:
            raw_text = extract_paragraph_text(child)
            text = normalize_whitespace(raw_text)
            if not is_empty_text(text):
                all_para_texts.append(text)
        elif child.tag == TAG_TBL:
            # 表格内的段落文本也收集
            table = DocxTable(child, docx)
            seen_cells = set()
            for row in table.rows:
                for cell in row.cells:
                    if cell._tc in seen_cells:
                        continue
                    seen_cells.add(cell._tc)
                    cell_text = extract_cell_text(cell)
                    if not is_empty_text(cell_text):
                        all_para_texts.append(cell_text)

    # 扫描文档中出现的章节类型
    chapter_types = scan_chapter_types(all_para_texts)

    # ---- 第二遍扫描：构建 blocks 和 anchors ----
    blocks: list[Block] = []
    anchors: list[Anchor] = []

    # 段落序号计数（基于全部段落，包含空段落）
    para_counter = 0
    # body 顶层元素顺序
    body_order = 0
    # block 全局阅读顺序
    block_order = 0
    # 表格编号
    table_index = 0
    # 列举组管理
    current_list_group_id: Optional[str] = None
    prev_block_for_list: Optional[Block] = None

    # 标题路径管理
    # 存储 (层级, 标题文本, heading_source) 的栈
    heading_stack: list[tuple[int, str, HeadingSource]] = []

    for child in body:
        if child.tag == TAG_P:
            # ---- 处理段落 ----
            raw_text = extract_paragraph_text(child)
            text = normalize_whitespace(raw_text)

            if is_empty_text(text):
                para_counter += 1  # 空段落不重排 para_index，但计数仍增加
                body_order += 1
                continue

            para_counter += 1
            style_name = _get_style_name(child, docx)

            # 检测编号
            has_numbering, numbering_text, numbering_unresolved = check_paragraph_numbering(
                child, docx, numbering_resolver
            )

            # 检测标题
            heading_result = detect_heading(text, style_name, chapter_types)
            if heading_result is not None:
                level, h_source = heading_result
                # 清理 heading_stack：弹出 >= 当前层级的标题
                while heading_stack and heading_stack[-1][0] >= level:
                    heading_stack.pop()
                heading_stack.append((level, text, h_source))
                # 当前 heading 的 section_path 是清理后的栈
                section_path = [h[1] for h in heading_stack]

                block = Block(
                    block_id=next_block_id(),
                    type=BlockType.HEADING,
                    text=text,
                    style=style_name,
                    section_path=section_path,
                    body_order=body_order,
                    block_order=block_order,
                    para_index=para_counter - 1,
                    table_index=None,
                    row_index=None,
                    cell_index=None,
                    has_numbering=has_numbering,
                    numbering_text=numbering_text,
                    numbering_unresolved=numbering_unresolved,
                    is_list_item=False,
                    list_group_id=None,
                    is_article_start=False,
                    heading_source=h_source,
                )
                blocks.append(block)
                body_order += 1
                block_order += 1

                # 生成 anchors
                _split_block_to_anchors(block, anchors, next_anchor_id)

                # 重置 list group
                current_list_group_id = None
                prev_block_for_list = block
                continue

            # 检测"第X条"
            is_art_start = is_article_start(text)

            # 检测列举项
            is_list = detect_chinese_list_item(text)

            # 非 heading 段落的 section_path
            section_path = [h[1] for h in heading_stack]

            # 列举组管理
            list_group_id: Optional[str] = None
            if is_list:
                if current_list_group_id is None:
                    # 检查前一个 non-empty paragraph 是否以：或:结尾
                    if prev_block_for_list is not None:
                        prev_text = prev_block_for_list.text.strip()
                        if prev_text.endswith("：") or prev_text.endswith(":"):
                            # 引导句也纳入同一 list group
                            if prev_block_for_list.list_group_id is None:
                                prev_block_for_list.list_group_id = next_list_group_id()
                            current_list_group_id = prev_block_for_list.list_group_id
                        else:
                            current_list_group_id = next_list_group_id()
                    else:
                        current_list_group_id = next_list_group_id()
                list_group_id = current_list_group_id
            else:
                current_list_group_id = None

            # 确定 block type
            if is_list:
                btype = BlockType.LIST_ITEM
            else:
                btype = BlockType.PARAGRAPH

            block = Block(
                block_id=next_block_id(),
                type=btype,
                text=text,
                style=style_name,
                section_path=section_path,
                body_order=body_order,
                block_order=block_order,
                para_index=para_counter - 1,
                table_index=None,
                row_index=None,
                cell_index=None,
                has_numbering=has_numbering,
                numbering_text=numbering_text,
                numbering_unresolved=numbering_unresolved,
                is_list_item=is_list,
                list_group_id=list_group_id,
                is_article_start=is_art_start,
                heading_source=None,
            )
            blocks.append(block)
            body_order += 1
            block_order += 1

            # 生成 anchors
            _split_block_to_anchors(block, anchors, next_anchor_id)

            prev_block_for_list = block

        elif child.tag == TAG_TBL:
            # ---- 处理表格 ----
            table = DocxTable(child, docx)
            tbl_body_order = body_order  # 表格作为一个顶层元素
            seen_cells = set()

            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # 合并单元格会由 python-docx 在 row.cells 中重复返回同一 w:tc。
                    if cell._tc in seen_cells:
                        continue
                    seen_cells.add(cell._tc)
                    cell_text = extract_cell_text(cell)
                    if is_empty_text(cell_text):
                        continue

                    section_path = [h[1] for h in heading_stack]

                    block = Block(
                        block_id=next_block_id(),
                        type=BlockType.TABLE_CELL,
                        text=cell_text,
                        style=None,
                        section_path=section_path,
                        body_order=tbl_body_order,  # 共享表格的 body_order
                        block_order=block_order,
                        para_index=None,  # table_cell 无 para_index
                        table_index=table_index,
                        row_index=row_idx,
                        cell_index=cell_idx,
                        has_numbering=False,
                        numbering_text=None,
                        numbering_unresolved=False,
                        is_list_item=False,
                        list_group_id=None,
                        is_article_start=False,
                        heading_source=None,
                    )
                    blocks.append(block)
                    block_order += 1

                    # 生成 anchors
                    _split_block_to_anchors(block, anchors, next_anchor_id)

            body_order += 1  # 表格算一个顶层元素
            table_index += 1
            current_list_group_id = None  # 表格中断列举组

        else:
            # 其他 body 元素（如 w:sectPr）按一个顶层元素计数
            body_order += 1

    # 脚注、尾注是法律文书中高密度的法源区域。按 Word note ID 和段落顺序
    # 追加到正文阅读流，保持 line 锚点全局连续且坐标可审计。
    for note_type, note_id, note_para_index, text in _extract_note_paragraphs(file_path):
        block = Block(
            block_id=next_block_id(),
            type=BlockType.FOOTNOTE if note_type == "footnote" else BlockType.ENDNOTE,
            text=text,
            style=None,
            section_path=["脚注" if note_type == "footnote" else "尾注"],
            body_order=body_order,
            block_order=block_order,
            para_index=note_para_index,
            note_type=note_type,
            note_id=note_id,
            has_numbering=False,
            numbering_unresolved=False,
            is_list_item=False,
            is_article_start=is_article_start(text),
        )
        blocks.append(block)
        _split_block_to_anchors(block, anchors, next_anchor_id)
        body_order += 1
        block_order += 1

    # ---- 构建 ParsedDocument ----
    doc_meta = DocMeta(
        schema_version="0.1",
        source_file=file_path,
        doc_hash=doc_hash,
    )

    parsed = ParsedDocument(
        doc_meta=doc_meta,
        blocks=blocks,
        anchors=anchors,
        chunks=[],  # 由 chunk_builder 后续填充
    )

    return parsed


def _extract_note_paragraphs(file_path: str) -> list[tuple[str, str, int, str]]:
    """从 DOCX 包内读取脚注/尾注正文，跳过 Word 的分隔符伪注释。"""
    notes: list[tuple[str, str, int, str]] = []
    with zipfile.ZipFile(file_path) as archive:
        for note_type, member, tag_name in (
            ("footnote", "word/footnotes.xml", "w:footnote"),
            ("endnote", "word/endnotes.xml", "w:endnote"),
        ):
            if member not in archive.namelist():
                continue
            root = parse_xml(archive.read(member))
            elements = list(root.findall(qn(tag_name)))
            elements.sort(key=lambda element: int(element.get(qn("w:id"), "0")))
            for element in elements:
                note_id = element.get(qn("w:id"), "")
                if note_id.lstrip("-").isdigit() and int(note_id) < 0:
                    continue
                for index, paragraph in enumerate(element.iter(qn("w:p"))):
                    text = normalize_whitespace(extract_paragraph_text(paragraph))
                    if text:
                        notes.append((note_type, note_id, index, text))
    return notes


def _get_style_name(para_element, docx: DocxDocument) -> Optional[str]:
    """
    获取段落的样式名称。

    Args:
        para_element: w:p XML 元素
        docx: python-docx Document 对象

    Returns:
        样式名称或 None
    """
    pPr = para_element.find(qn("w:pPr"))
    if pPr is None:
        return None
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        return None
    style_id = pStyle.get(qn("w:val"))
    if style_id is None:
        return None
    style = next((item for item in docx.styles if item.style_id == style_id), None)
    return style.name if style is not None else style_id


def _split_block_to_anchors(
    block: Block,
    anchors: list[Anchor],
    next_anchor_id,
):
    """
    为 block 文本分句，生成 anchors 并更新 block 的 anchor 关联字段。

    Args:
        block: Block 对象（会被原地修改）
        anchors: 全局 anchors 列表（会被原地追加）
        next_anchor_id: ID 生成器函数
    """
    sentences = split_sentences(block.text)
    if not sentences:
        return

    sentence_anchors: list[str] = []
    for sent in sentences:
        anchor = Anchor(
            anchor=next_anchor_id(),
            text=sent.text,
            block_id=block.block_id,
            para_index=block.para_index,
            note_type=block.note_type,
            note_id=block.note_id,
            char_start=sent.char_start,
            char_end=sent.char_end,
        )
        anchors.append(anchor)
        sentence_anchors.append(anchor.anchor)

    # 更新 block 的 anchor 关联
    block.sentence_anchors = sentence_anchors
    block.anchor_range = [sentence_anchors[0], sentence_anchors[-1]]
