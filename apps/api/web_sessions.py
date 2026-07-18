"""网页版短期文档会话与安全 DOCX 修订。"""

from __future__ import annotations

import io
import threading
from dataclasses import dataclass, field
from datetime import datetime, timedelta, timezone
from uuid import uuid4

from docx import Document

from ccitecheck.domain.result import FrontendVerificationDocument
from ccitecheck.domain.revisions import RevisionProposal

SESSION_TTL = timedelta(hours=1)


@dataclass
class WebDocumentSession:
    session_id: str
    file_name: str
    document_bytes: bytes
    revisions: dict[str, RevisionProposal]
    expires_at: datetime
    accepted: set[str] = field(default_factory=set)


class WebSessionStore:
    def __init__(self) -> None:
        self._sessions: dict[str, WebDocumentSession] = {}
        self._lock = threading.Lock()

    def create(self, file_name: str, document_bytes: bytes, verification: FrontendVerificationDocument) -> WebDocumentSession:
        now = datetime.now(timezone.utc)
        session = WebDocumentSession(
            session_id=uuid4().hex,
            file_name=file_name,
            document_bytes=document_bytes,
            revisions=_collect_revisions(verification),
            expires_at=now + SESSION_TTL,
        )
        with self._lock:
            self._purge(now)
            self._sessions[session.session_id] = session
        timer = threading.Timer(SESSION_TTL.total_seconds(), self.delete, args=(session.session_id,))
        timer.daemon = True
        timer.start()
        return session

    def delete(self, session_id: str) -> None:
        with self._lock:
            self._sessions.pop(session_id, None)

    def get(self, session_id: str) -> WebDocumentSession | None:
        now = datetime.now(timezone.utc)
        with self._lock:
            self._purge(now)
            return self._sessions.get(session_id)

    def set_revision(self, session_id: str, check_id: str, accepted: bool) -> WebDocumentSession | None:
        with self._lock:
            session = self._sessions.get(session_id)
            if not session or session.expires_at <= datetime.now(timezone.utc):
                self._sessions.pop(session_id, None)
                return None
            if check_id not in session.revisions:
                raise KeyError(check_id)
            if accepted:
                session.accepted.add(check_id)
            else:
                session.accepted.discard(check_id)
            return session

    def _purge(self, now: datetime) -> None:
        expired = [key for key, value in self._sessions.items() if value.expires_at <= now]
        for key in expired:
            del self._sessions[key]


def render_revised_docx(session: WebDocumentSession) -> bytes:
    document = Document(io.BytesIO(session.document_bytes))
    for check_id in sorted(session.accepted):
        proposal = session.revisions[check_id]
        if not proposal.machine_applicable or not proposal.revised_text:
            continue
        matches = [paragraph for paragraph in _paragraphs(document) if paragraph.text.count(proposal.original_text)]
        if sum(paragraph.text.count(proposal.original_text) for paragraph in matches) != 1:
            raise ValueError(f"修订 {check_id} 的原文已无法唯一定位")
        _replace_in_paragraph(matches[0], proposal.original_text, proposal.revised_text)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _collect_revisions(verification: FrontendVerificationDocument) -> dict[str, RevisionProposal]:
    revisions: dict[str, RevisionProposal] = {}
    for check in [*verification.statute_results, *verification.case_results]:
        proposals = [finding.revision for finding in check.findings if finding.revision and finding.revision.machine_applicable]
        if len(proposals) == 1:
            revisions[check.check_id] = proposals[0]
    return revisions


def _paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _replace_in_paragraph(paragraph, original: str, revised: str) -> None:
    full_text = paragraph.text
    start = full_text.index(original)
    end = start + len(original)
    positions: list[tuple[object, int, int]] = []
    cursor = 0
    for run in paragraph.runs:
        positions.append((run, cursor, cursor + len(run.text)))
        cursor += len(run.text)
    first = next((run for run, left, right in positions if left <= start < right), None)
    if first is None:
        raise ValueError("修订原文无法映射到 Word 文本片段")
    for run, left, right in positions:
        if right <= start or left >= end:
            continue
        before = run.text[: max(0, start - left)] if left <= start < right else ""
        after = run.text[max(0, end - left):] if left < end <= right else ""
        run.text = before + (revised if run is first else "") + after


WEB_SESSIONS = WebSessionStore()

__all__ = ["WEB_SESSIONS", "render_revised_docx"]
