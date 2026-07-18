"""为 Word、飞书和其他客户端提供统一核查能力的 FastAPI 应用。"""

from __future__ import annotations

import base64
import binascii
import hashlib
import io
import os
import tempfile
import time
from collections import defaultdict, deque
from pathlib import Path
from urllib.parse import quote

from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, RedirectResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles

from ccitecheck.application import (
    DocumentPipelineError,
    extract_document_claims,
    parse_and_validate_document,
    validate_shared_document,
    verify_document_claims,
)
from ccitecheck.infrastructure.paths import PROJECT_ROOT
from ccitecheck.infrastructure.config import load_project_env
from ccitecheck.parsing.feishu import parse_feishu_snapshot
from ccitecheck.output import summarize_verification

from .schema import (
    DebugEventRequest,
    DocumentCheckRequest,
    DocumentCheckResponse,
    FeishuDocumentCheckRequest,
    SelectionCheckRequest,
    WebCheckResponse,
    WebPreviewBlock,
    WebRevisionRequest,
    WebTextCheckRequest,
)
from .debug_capture import append_event, create_run, write_json
from .web_sessions import WEB_SESSIONS, render_revised_docx

ADDIN_ROOT = PROJECT_ROOT / "apps" / "word_addin"
FEISHU_ADDIN_ROOT = PROJECT_ROOT / "apps" / "feishu"
WEB_ROOT = PROJECT_ROOT / "apps" / "web"
LAW_DB = PROJECT_ROOT / "data" / "laws.sqlite"
MAX_DOCUMENT_BYTES = 25 * 1024 * 1024

load_project_env()
app = FastAPI(title="CCiteheck API", version="1.0.0")
_web_requests: dict[str, deque[float]] = defaultdict(deque)
allowed_origins = [
    origin.strip()
    for origin in os.getenv("CCITECHECK_ALLOWED_ORIGINS", "").split(",")
    if origin.strip()
]
if allowed_origins:
    app.add_middleware(
        CORSMiddleware,
        allow_origins=allowed_origins,
        allow_methods=["POST", "OPTIONS"],
        allow_headers=["Content-Type"],
    )
@app.middleware("http")
async def limit_public_checks(request: Request, call_next):
    if request.method == "POST" and request.url.path in {"/api/web/checks", "/api/web/checks/text"}:
        client = request.client.host if request.client else "unknown"
        now = time.monotonic()
        recent = _web_requests[client]
        while recent and recent[0] < now - 600:
            recent.popleft()
        if len(recent) >= 6:
            return JSONResponse(status_code=429, content={"detail": "核查请求过于频繁，请十分钟后再试"})
        recent.append(now)
    return await call_next(request)


@app.middleware("http")
async def revalidate_static_assets(request, call_next):
    """Office WebView 磁盘缓存极顽固；静态资源强制每次向服务端校验新鲜度。"""
    response = await call_next(request)
    path = request.url.path
    if path.startswith(("/assets", "/web")) or path.endswith((".html", ".css", ".js")):
        response.headers["Cache-Control"] = "no-cache"
    if path == "/" or path.startswith("/web"):
        response.headers["Content-Security-Policy"] = (
            "default-src 'self'; script-src 'self'; style-src 'self'; "
            "connect-src 'self'; img-src 'self' data:; frame-ancestors 'none'"
        )
        response.headers["Referrer-Policy"] = "no-referrer"
        response.headers["X-Content-Type-Options"] = "nosniff"
    return response


app.mount("/assets", StaticFiles(directory=ADDIN_ROOT / "assets"), name="assets")
app.mount(
    "/feishu-addon",
    StaticFiles(directory=FEISHU_ADDIN_ROOT, html=True),
    name="feishu-addon",
)
app.mount("/web-assets", StaticFiles(directory=WEB_ROOT / "assets"), name="web-assets")


@app.get("/", include_in_schema=False)
def root() -> RedirectResponse:
    return RedirectResponse("/web/")


@app.get("/web/", include_in_schema=False)
def web_app() -> FileResponse:
    return FileResponse(WEB_ROOT / "index.html")


@app.get("/taskpane.html", include_in_schema=False)
def taskpane() -> FileResponse:
    return FileResponse(ADDIN_ROOT / "taskpane.html")


@app.get("/help.html", include_in_schema=False)
def help_page() -> FileResponse:
    return FileResponse(ADDIN_ROOT / "help.html")


@app.get("/api/health")
def health() -> dict[str, str | bool]:
    return {
        "status": "ok",
        "pkulaw_configured": bool(os.getenv("PKULAW_ACCESS_TOKEN", "").strip()),
        "llm_configured": bool(
            (os.getenv("DASHSCOPE_API_KEY") or os.getenv("LLM_API_KEY") or "").strip()
        ),
    }


def _validate_scope(request) -> None:
    if not (request.include_statutes or request.include_cases):
        raise HTTPException(status_code=400, detail="请至少选择一种核查范围（法规引用或司法案例）")


@app.post("/api/checks", response_model=DocumentCheckResponse)
def check_document(request: DocumentCheckRequest) -> DocumentCheckResponse:
    return _run_document_check(request, capture_debug=True)


def _run_document_check(request: DocumentCheckRequest, *, capture_debug: bool) -> DocumentCheckResponse:
    _validate_scope(request)
    document_bytes = _decode_document(request.docx_base64)
    debug_run_id = create_run("document", document_bytes) if capture_debug else None
    write_json(debug_run_id, "request.json", {
        **request.model_dump(exclude={"docx_base64"}),
        "docx_base64_length": len(request.docx_base64),
    })
    if len(document_bytes) > MAX_DOCUMENT_BYTES:
        raise HTTPException(status_code=413, detail="文档超过 25 MB 限制")
    if not document_bytes.startswith(b"PK"):
        raise HTTPException(status_code=400, detail="文件不是有效的 DOCX 文档")

    try:
        with tempfile.TemporaryDirectory(prefix="ccitecheck-document-") as temporary_dir:
            document_path = Path(temporary_dir) / "document.docx"
            document_path.write_bytes(document_bytes)
            parsed_document = parse_and_validate_document(document_path)
            write_json(debug_run_id, "parsed-document.json", parsed_document)
            claim_document = extract_document_claims(
                parsed_document,
                include_statutes=request.include_statutes,
                include_cases=request.include_cases,
            )
            write_json(debug_run_id, "claim-document.json", claim_document)
            verification = verify_document_claims(
                claim_document,
                LAW_DB,
                semantic_check=request.semantic_check,
                include_statutes=request.include_statutes,
                include_cases=request.include_cases,
            )
    except DocumentPipelineError as exc:
        write_json(debug_run_id, "error.json", {"type": type(exc).__name__, "message": str(exc)})
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except Exception as exc:
        write_json(debug_run_id, "error.json", {"type": type(exc).__name__, "message": str(exc)})
        raise

    response = DocumentCheckResponse(
        file_name=Path(request.file_name).name,
        document_key="sha256:" + hashlib.sha256(document_bytes).hexdigest(),
        semantic_check=request.semantic_check,
        summary=summarize_verification(verification),
        verification=verification,
        debug_run_id=debug_run_id,
    )
    write_json(debug_run_id, "response.json", response)
    return response


@app.post("/api/web/checks", response_model=WebCheckResponse)
def check_web_document(request: DocumentCheckRequest) -> WebCheckResponse:
    document_bytes = _decode_document(request.docx_base64)
    # 公网网页文书不进入本地 debug_runs，确保一小时会话是唯一临时副本。
    response = _run_document_check(request, capture_debug=False)
    parsed = _parse_preview(document_bytes)
    session = WEB_SESSIONS.create(response.file_name, document_bytes, response.verification)
    return WebCheckResponse(
        **response.model_dump(),
        session_id=session.session_id,
        expires_at=session.expires_at.isoformat(),
        preview_blocks=parsed,
    )


@app.post("/api/web/checks/text", response_model=WebCheckResponse)
def check_web_text(request: WebTextCheckRequest) -> WebCheckResponse:
    from docx import Document as DocxDocument

    document = DocxDocument()
    for line in request.text.splitlines():
        document.add_paragraph(line)
    output = io.BytesIO()
    document.save(output)
    return check_web_document(DocumentCheckRequest(
        file_name=request.file_name,
        docx_base64=base64.b64encode(output.getvalue()).decode("ascii"),
        semantic_check=request.semantic_check,
        include_statutes=request.include_statutes,
        include_cases=request.include_cases,
    ))


@app.post("/api/web/sessions/{session_id}/revisions")
def accept_web_revision(session_id: str, request: WebRevisionRequest) -> dict:
    try:
        session = WEB_SESSIONS.set_revision(session_id, request.check_id, True)
    except KeyError as exc:
        raise HTTPException(status_code=400, detail="该核查项没有可自动应用的安全修订") from exc
    if not session:
        raise HTTPException(status_code=404, detail="核查会话已过期，请重新上传文书")
    return {"accepted_check_ids": sorted(session.accepted)}


@app.delete("/api/web/sessions/{session_id}/revisions/{check_id}")
def undo_web_revision(session_id: str, check_id: str) -> dict:
    try:
        session = WEB_SESSIONS.set_revision(session_id, check_id, False)
    except KeyError as exc:
        raise HTTPException(status_code=400, detail="该核查项没有可撤销的安全修订") from exc
    if not session:
        raise HTTPException(status_code=404, detail="核查会话已过期，请重新上传文书")
    return {"accepted_check_ids": sorted(session.accepted)}


@app.get("/api/web/sessions/{session_id}/document")
def download_web_document(session_id: str):
    session = WEB_SESSIONS.get(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="核查会话已过期，请重新上传文书")
    try:
        content = render_revised_docx(session)
    except ValueError as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc
    stem = Path(session.file_name).stem
    filename = f"{stem}_CCiteheck修订版.docx"
    return StreamingResponse(
        io.BytesIO(content),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"},
    )


def _parse_preview(document_bytes: bytes) -> list[WebPreviewBlock]:
    with tempfile.TemporaryDirectory(prefix="ccitecheck-preview-") as temporary_dir:
        path = Path(temporary_dir) / "document.docx"
        path.write_bytes(document_bytes)
        parsed = parse_and_validate_document(path)
    return [
        WebPreviewBlock(
            block_id=block.external_block_id or block.block_id,
            block_type=block.type.value,
            text=block.text,
            order=block.block_order,
        )
        for block in parsed.blocks
        if block.text.strip()
    ]


@app.post("/api/checks/selection", response_model=DocumentCheckResponse)
def check_selection(request: SelectionCheckRequest) -> DocumentCheckResponse:
    """核查用户在 Word 中选中的文本片段：构造临时 DOCX 复用完整核查管线。"""
    _validate_scope(request)
    lines = [line.strip() for line in request.text.splitlines() if line.strip()]
    if not lines:
        raise HTTPException(status_code=400, detail="选中内容为空，无法核查")

    from docx import Document as DocxDocument

    debug_document = (
        _decode_document(request.debug_docx_base64)
        if request.debug_docx_base64
        else None
    )
    debug_run_id = create_run("selection", debug_document)
    write_json(debug_run_id, "request.json", {
        **request.model_dump(exclude={"debug_docx_base64"}),
        "debug_docx_base64_length": len(request.debug_docx_base64 or ""),
    })

    try:
        with tempfile.TemporaryDirectory(prefix="ccitecheck-selection-") as temporary_dir:
            selection_path = Path(temporary_dir) / "selection.docx"
            selection_doc = DocxDocument()
            for line in lines:
                selection_doc.add_paragraph(line)
            selection_doc.save(selection_path)
            parsed_document = parse_and_validate_document(selection_path)
            write_json(debug_run_id, "parsed-selection.json", parsed_document)
            claim_document = extract_document_claims(
                parsed_document,
                include_statutes=request.include_statutes,
                include_cases=request.include_cases,
            )
            write_json(debug_run_id, "claim-document.json", claim_document)
            verification = verify_document_claims(
                claim_document,
                LAW_DB,
                semantic_check=request.semantic_check,
                include_statutes=request.include_statutes,
                include_cases=request.include_cases,
            )
            _rebase_selection_locations(verification, request.source_blocks)
    except DocumentPipelineError as exc:
        write_json(debug_run_id, "error.json", {"type": type(exc).__name__, "message": str(exc)})
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except Exception as exc:
        write_json(debug_run_id, "error.json", {"type": type(exc).__name__, "message": str(exc)})
        raise

    response = DocumentCheckResponse(
        file_name=f"{Path(request.file_name).name}（选中片段）",
        document_key="sha256:" + hashlib.sha256(
            (Path(request.file_name).name + "\0" + request.text).encode("utf-8")
        ).hexdigest(),
        semantic_check=request.semantic_check,
        summary=summarize_verification(verification),
        verification=verification,
        debug_run_id=debug_run_id,
    )
    write_json(debug_run_id, "response.json", response)
    return response


def _rebase_selection_locations(verification, source_blocks) -> None:
    """把临时选区 DOCX 的段落坐标映射回当前 Word 文档。"""
    located_items = [*verification.statute_results, *verification.case_results]
    for check in located_items:
        rebased = []
        for location in check.source_locations:
            parts = location.block_id.split(":")
            if len(parts) != 3 or parts[:2] != ["word", "p"]:
                continue
            index = int(parts[2])
            if index >= len(source_blocks):
                continue
            source = source_blocks[index]
            location.block_id = source.block_id
            location.char_start += source.char_start
            location.char_end += source.char_start
            # 选区临时 DOCX 中的出现序号不等于原文所属段落中的出现序号。
            location.occurrence = None
            rebased.append(location)
        check.source_locations = rebased


@app.post("/api/debug-events", status_code=204)
def capture_debug_event(request: DebugEventRequest) -> None:
    try:
        append_event(request.run_id, {
            "event": request.event,
            "payload": request.payload,
        })
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc


@app.post("/api/feishu/checks", response_model=DocumentCheckResponse)
def check_feishu_document(request: FeishuDocumentCheckRequest) -> DocumentCheckResponse:
    """让飞书块快照复用与 Word 相同的识别、溯源和判定流水线。"""
    _validate_scope(request)
    try:
        parsed_document = validate_shared_document(
            parse_feishu_snapshot(request.snapshot)
        )
        claim_document = extract_document_claims(
            parsed_document,
            include_statutes=request.include_statutes,
            include_cases=request.include_cases,
        )
        verification = verify_document_claims(
            claim_document,
            LAW_DB,
            semantic_check=request.semantic_check,
            include_statutes=request.include_statutes,
            include_cases=request.include_cases,
        )
    except DocumentPipelineError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc

    return DocumentCheckResponse(
        file_name=request.snapshot.title,
        document_key=parsed_document.doc_meta.doc_hash,
        semantic_check=request.semantic_check,
        summary=summarize_verification(verification),
        verification=verification,
    )


def _decode_document(encoded: str) -> bytes:
    try:
        return base64.b64decode(encoded, validate=True)
    except (binascii.Error, ValueError) as exc:
        raise HTTPException(status_code=400, detail="DOCX Base64 数据无效") from exc
