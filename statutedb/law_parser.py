"""
法规文本解析：官方 Word(.docx)/纯文本 → StatuteDoc 结构。

预期输入是国家法律法规数据库（flk.npc.gov.cn）的官方 Word 下载，
兼容从任意来源复制保存的纯文本。

结构识别规则：
  - 首个非空行为标题（可被调用方覆盖）
  - 标题后的（…）行提取颁布信息：文号、通过/修正日期、施行日期
  - ^第X编/分编/章/节 → 层级标题（维护 section_path 栈）
  - ^第X条(之X)? → 新条文，同行余文为第一款起始
  - 条内每个段落为一款；（一）开头的段落为项，归属前一款
  - 条号连续性校验：跳号/回退记入 warnings（防解析错位），不中断
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from pydantic import BaseModel, Field

from .cn_num import cn_to_int
from .normalizer import normalize_title, strip_parenthetical

_CN_NUM = r"[零一二两三四五六七八九十百千万0-9]+"

# 条文起始：行首"第X条"或"第X条之Y"，其后是分隔空白或直接接正文
_ARTICLE_START_PATTERN = re.compile(
    rf"^第({_CN_NUM})条(?:之({_CN_NUM}))?(?![零一二两三四五六七八九十百千万0-9])[　\s]*"
)

# 层级标题：第X编/分编/章/节（标题行不应过长，防止误吞正文）
_SECTION_PATTERN = re.compile(
    rf"^(第{_CN_NUM}(?:编|分编|章|节))[　\s]*(.*)$"
)
_SECTION_LEVELS = {"编": 0, "分编": 1, "章": 2, "节": 3}

# 司法解释常用编号章节标题："一、一般规定"（短、无句读、顿号分隔）
# 约束从紧：条文内的款极少以"X、"开头（列举用（一）格式），
# 且款必然是含句读的完整句子
_NUMBERED_SECTION_PATTERN = re.compile(
    r"^([一二三四五六七八九十]+)、([^。，；：、]{1,20})$"
)

# 项起始：（一）/ (一)
_ITEM_PATTERN = re.compile(r"^[（(]([零一二两三四五六七八九十0-9]+)[）)]")

# 文号：法释〔2022〕6号 / 国令第752号 等
_DOC_NUMBER_PATTERN = re.compile(r"(法释〔\d{4}〕\d+号|法发〔\d{4}〕\d+号|国务院令第\d+号)")

_DATE_PATTERN = re.compile(r"(\d{4})年(\d{1,2})月(\d{1,2})日")
_EFFECTIVE_PATTERN = re.compile(r"自(\d{4})年(\d{1,2})月(\d{1,2})日起施行")


class StatuteParagraph(BaseModel):
    """一款；items 为该款下的各项原文（含（一）编号）。"""
    text: str
    items: list[str] = Field(default_factory=list)


class StatuteArticle(BaseModel):
    """一条。"""
    article_num: int
    article_suffix: int = 0
    article_label: str = Field(description="原文条号标签，如'第一百八十四条'")
    section_path: str = Field(default="", description="'第一编 总则/第八章 …'")
    paragraphs: list[StatuteParagraph] = Field(default_factory=list)

    @property
    def full_text(self) -> str:
        """全条原文：条号 + 各款（款间换行，项各占一行）。"""
        lines: list[str] = []
        for i, para in enumerate(self.paragraphs):
            prefix = f"{self.article_label}　" if i == 0 else ""
            lines.append(prefix + para.text)
            lines.extend(para.items)
        return "\n".join(lines)


class StatuteDoc(BaseModel):
    """解析后的一部法规。"""
    title: str
    doc_number: Optional[str] = None
    promulgated_on: Optional[str] = None
    effective_on: Optional[str] = None
    version_note: Optional[str] = None
    articles: list[StatuteArticle] = Field(default_factory=list)
    warnings: list[str] = Field(default_factory=list)


def parse_law_file(path: str, title_override: Optional[str] = None) -> StatuteDoc:
    """解析 .docx 或纯文本法规文件。"""
    p = Path(path)
    if p.suffix.lower() == ".docx":
        lines = _read_docx_lines(p)
    else:
        lines = p.read_text(encoding="utf-8").splitlines()
    return parse_law_lines(lines, title_override=title_override)


def _read_docx_lines(path: Path) -> list[str]:
    from docx import Document

    doc = Document(str(path))
    return [para.text for para in doc.paragraphs]


def parse_law_lines(
    lines: list[str],
    title_override: Optional[str] = None,
) -> StatuteDoc:
    """
    核心解析逻辑（与文件格式解耦，便于测试）。

    Args:
        lines: 逐段文本（docx 段落或 txt 行）
        title_override: 显式指定标题时跳过首行推断
    """
    stripped = [ln.strip() for ln in lines]
    non_empty = [(i, ln) for i, ln in enumerate(stripped) if ln]
    if not non_empty:
        raise ValueError("empty law document")

    # ---- 标题 ----
    if title_override:
        title = normalize_title(title_override)
        body_start = 0
    else:
        first_idx, first_line = non_empty[0]
        title = normalize_title(first_line)
        body_start = first_idx + 1

    doc = StatuteDoc(title=title)

    # 标题括号注记 → version_note
    base_title = strip_parenthetical(title)
    if base_title != title:
        doc.version_note = title[len(base_title):].strip("（）()")
        doc.title = title  # 保留全称原样

    # ---- 逐行状态机 ----
    section_stack: list[tuple[int, str]] = []   # (level, "第一编 总则")
    current_article: Optional[StatuteArticle] = None
    metadata_lines: list[str] = []
    seen_first_article = False

    for line in stripped[body_start:]:
        if not line:
            continue

        # 条文起始
        m = _ARTICLE_START_PATTERN.match(line)
        if m:
            article = _start_article(m, line, section_stack, doc)
            if article is not None:
                current_article = article
                seen_first_article = True
                continue
            # 条号解析失败：当普通款处理（记入 warnings）

        # 层级标题（仅在标题行短促时认定，防止吞正文；目录区同规则处理，
        # 正文中同名标题会重置栈，目录残留不影响条文归属）
        sm = _SECTION_PATTERN.match(line)
        if sm and _looks_like_heading(line, sm):
            _push_section(section_stack, sm)
            continue

        # 司法解释编号章节："一、一般规定"（视同"章"级）
        nm = _NUMBERED_SECTION_PATTERN.match(line)
        if nm:
            _push_numbered_section(section_stack, line)
            continue

        if current_article is None:
            # 条文开始前的散文：颁布信息等元数据
            if not seen_first_article:
                metadata_lines.append(line)
            continue

        # 条内内容：项 或 新款
        if _ITEM_PATTERN.match(line):
            if current_article.paragraphs:
                current_article.paragraphs[-1].items.append(line)
            else:
                # 条文首行即是项（异常格式）：并入第一款
                current_article.paragraphs.append(
                    StatuteParagraph(text=line)
                )
                doc.warnings.append(
                    f"{current_article.article_label}: 首款以项编号开头"
                )
        else:
            current_article.paragraphs.append(StatuteParagraph(text=line))

    _extract_metadata("\n".join(metadata_lines), doc)
    _check_continuity(doc)
    return doc


def _start_article(
    m: re.Match,
    line: str,
    section_stack: list[tuple[int, str]],
    doc: StatuteDoc,
) -> Optional[StatuteArticle]:
    """从匹配行构建新条文；条号不可解析时返回 None。"""
    try:
        num = cn_to_int(m.group(1))
        suffix = cn_to_int(m.group(2)) if m.group(2) else 0
    except ValueError:
        doc.warnings.append(f"条号解析失败: {line[:30]}")
        return None

    label = f"第{m.group(1)}条" + (f"之{m.group(2)}" if m.group(2) else "")
    article = StatuteArticle(
        article_num=num,
        article_suffix=suffix,
        article_label=label,
        section_path="/".join(name for _, name in section_stack),
    )
    rest = line[m.end():].strip()
    if rest:
        article.paragraphs.append(StatuteParagraph(text=rest))
    doc.articles.append(article)
    return article


def _looks_like_heading(line: str, m: re.Match) -> bool:
    """
    层级标题判定：标题名部分较短且不含句读。

    排除"第一章规定的…"这类正文引用（标题行没有句号/逗号，
    且总长通常 < 30 字）。
    """
    rest = m.group(2)
    if len(line) > 30:
        return False
    if any(punct in rest for punct in "。，；：、"):
        return False
    return True


def _push_section(section_stack: list[tuple[int, str]], m: re.Match) -> None:
    """按层级维护 section 栈：同级或更深的旧标题出栈。"""
    marker = m.group(1)
    for unit, level in _SECTION_LEVELS.items():
        if marker.endswith(unit):
            break
    else:
        return
    # 官方文本标题名常含全角空格（"总　　则"），压缩为无空格
    name = (marker + " " + re.sub(r"[\s　]+", "", m.group(2))).strip()
    while section_stack and section_stack[-1][0] >= level:
        section_stack.pop()
    section_stack.append((level, name))


def _push_numbered_section(
    section_stack: list[tuple[int, str]], line: str
) -> None:
    """"一、一般规定"式标题按"章"级（level=2）入栈。"""
    level = _SECTION_LEVELS["章"]
    name = re.sub(r"[\s　]+", "", line)
    while section_stack and section_stack[-1][0] >= level:
        section_stack.pop()
    section_stack.append((level, name))


def _extract_metadata(meta_text: str, doc: StatuteDoc) -> None:
    """从条文前的散文提取文号与日期。"""
    if not meta_text:
        return
    dn = _DOC_NUMBER_PATTERN.search(meta_text)
    if dn:
        doc.doc_number = dn.group(1)
    eff = _EFFECTIVE_PATTERN.search(meta_text)
    if eff:
        doc.effective_on = f"{eff.group(1)}-{int(eff.group(2)):02d}-{int(eff.group(3)):02d}"
    first_date = _DATE_PATTERN.search(meta_text)
    if first_date:
        doc.promulgated_on = (
            f"{first_date.group(1)}-{int(first_date.group(2)):02d}"
            f"-{int(first_date.group(3)):02d}"
        )


def _check_continuity(doc: StatuteDoc) -> None:
    """
    条号连续性校验：条号应单调递增（允许"之X"插入条）。

    跳号或回退说明解析可能错位（如把正文误认成条文起始），
    记入 warnings 供导入时人工复核。
    """
    prev: Optional[tuple[int, int]] = None
    for art in doc.articles:
        cur = (art.article_num, art.article_suffix)
        if prev is not None:
            expected_next = [
                (prev[0] + 1, 0),           # 正常递增
                (prev[0], prev[1] + 1),     # 之X 插入条
            ]
            if cur not in expected_next and cur <= prev:
                doc.warnings.append(
                    f"条号非递增: {art.article_label} 出现在第{prev[0]}条"
                    f"{'之' + str(prev[1]) if prev[1] else ''}之后"
                )
            elif cur not in expected_next:
                doc.warnings.append(
                    f"条号跳跃: 第{prev[0]}条后出现{art.article_label}"
                )
        prev = cur

    if not doc.articles:
        doc.warnings.append("未解析到任何条文")
